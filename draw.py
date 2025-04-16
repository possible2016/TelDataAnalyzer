from __future__ import unicode_literals
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import pandas as pd
import matplotlib
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import mplcursors
import numpy as np
import datetime
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import json
import matplotlib.dates as mdates
from matplotlib.ticker import MaxNLocator



matplotlib.use('TkAgg')

# 设置字体以支持中文
matplotlib.rcParams['font.sans-serif'] = ['SimHei']  # 使用黑体
matplotlib.rcParams['axes.unicode_minus'] = False  # 正常显示负号


def select_folder():
    folder_path = filedialog.askdirectory(title="选择文件夹")
    return folder_path


def select_files(initialdir):
    filetypes = [("Excel文件", "*.xls *.xlsx")]
    filenames = filedialog.askopenfilenames(title="选择文件", initialdir=initialdir, filetypes=filetypes)
    return filenames


def read_data(filename):
    try:
        if filename.endswith('.xlsx'):
            df = pd.read_excel(filename, engine='openpyxl')
        elif filename.endswith('.xls'):
            # 尝试使用不同的编码读取文件
            try:
                with open(filename, "rb") as file1:
                    data = file1.read()
                data = data.decode('utf-8').splitlines()
            except UnicodeDecodeError:
                try:
                    with open(filename, "rb") as file1:
                        data = file1.read()
                    data = data.decode('ANSI').splitlines()
                except UnicodeDecodeError:
                    with open(filename, "rb") as file1:
                        data = file1.read()
                    data = data.decode('gbk').splitlines()
            # 将数据转换为DataFrame
            df = pd.DataFrame([row.split('\t') for row in data])  # 使用制表符分隔符
            df.columns = df.iloc[0].astype(str)  # 将第一行设置为列名，并确保列名是字符串类型
            df = df[1:]  # 删除第一行
        else:
            messagebox.showerror("错误", "文件格式不支持")
            return None

        # 打印DataFrame以进行调试
        # print(f"读取文件: {filename}")
        # print(df.head())

        # 如果列数超过256，分块读取
        if df.shape[1] > 256:
            chunks = [df.iloc[:, i:i+256] for i in range(0, df.shape[1], 256)]
            df = pd.concat(chunks, axis=1)
    except Exception as e:
        messagebox.showerror("错误", str(e))
        return None
    return df


def clean_data(series):
    return pd.to_numeric(series, errors='coerce')
def create_main_interface(folder_path):
    main_window = tk.Tk()
    main_window.title("参数折线图")
    main_window.geometry("1200x800")
    chart_type_var = tk.StringVar(value="柱状图")

    files = []
    file_names = []
    current_df = None
    y_axis_vars = {}

    file_var = tk.StringVar()
    file_combobox = ttk.Combobox(main_window, textvariable=file_var, state='readonly')
    file_combobox.place(x=200, y=10, width=500, height=20)

    def add_files():
        nonlocal current_df
        new_files = select_files(folder_path)
        if new_files:
            files.extend(new_files)
            file_names.extend([os.path.basename(f) for f in new_files])
            file_combobox['values'] = file_names
            if not file_var.get():
                file_var.set(file_names[0])
            current_df = read_data(files[0])
            update_column_selection(current_df)

    add_button = tk.Button(main_window, text="选择文件", command=add_files, font=("Arial", 12))
    add_button.place(x=700, y=10, width=100, height=20)

    column_frame = tk.Frame(main_window)
    column_frame.place(x=800, y=60, width=350, height=400)

    x_axis_var = tk.StringVar()

    def update_column_selection(df):
        """更新列名选择框"""
        nonlocal y_axis_vars
        for widget in column_frame.winfo_children():
            widget.destroy()

        if df is None:
            return

        # 横坐标选择（单选下拉框）
        tk.Label(column_frame, text="选择横坐标（单选）:", font=("Arial", 12)).pack(anchor='w')
        x_axis_combobox = ttk.Combobox(column_frame, textvariable=x_axis_var, values=list(df.columns), state='readonly')
        x_axis_combobox.pack(anchor='w', fill='x')
        if df.columns.any():
            x_axis_var.set(df.columns[0])  # 默认选择第一列

        # 纵坐标选择（多选复选框 + 滚动条）
        tk.Label(column_frame, text="选择纵坐标（多选）:", font=("Arial", 12)).pack(anchor='w')

        # 创建一个带有边框的Frame来容纳Canvas和Scrollbar
        frame_with_border = tk.Frame(column_frame, relief="solid", bd=0.5)  # 添加黑色边框
        frame_with_border.pack(fill="both", expand=True)

        # 创建Canvas
        scroll_canvas = tk.Canvas(frame_with_border)
        scrollbar = tk.Scrollbar(frame_with_border, orient="vertical", command=scroll_canvas.yview)

        # 绑定Scrollbar与Canvas的同步滚动
        scroll_canvas.configure(yscrollcommand=scrollbar.set)

        # 创建内嵌Frame用于存放Checkbuttons
        frame_inside_canvas = tk.Frame(scroll_canvas)

        # 将Frame放入Canvas内部
        window_inside_canvas = scroll_canvas.create_window((0, 0), window=frame_inside_canvas, anchor="nw")

        def on_mouse_wheel(event):
            """当鼠标在该区域时滚动"""
            if frame_with_border.winfo_containing(event.x_root, event.y_root):
                scroll_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        # 绑定鼠标滚轮事件，仅在鼠标位于该区域时触发
        scroll_canvas.bind("<Enter>", lambda e: scroll_canvas.bind_all("<MouseWheel>", on_mouse_wheel))
        scroll_canvas.bind("<Leave>", lambda e: scroll_canvas.unbind_all("<MouseWheel>"))

        frame_inside_canvas.bind(
            "<Configure>",
            lambda e: scroll_canvas.configure(
                scrollregion=scroll_canvas.bbox("all")
            )
        )

        # 创建纵坐标复选框
        y_axis_vars = {}
        for col in df.columns:
            var = tk.IntVar()
            chk = tk.Checkbutton(frame_inside_canvas, text=col, variable=var)
            chk.pack(anchor='w', padx=5, pady=2)
            y_axis_vars[col] = var

        # **将Canvas和Scrollbar放入容器中**
        scroll_canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")  # **滚动条可见且可拖动**

    def collect_selected_columns():
        """收集用户选择的列"""
        selected_columns = [col for col, var in y_axis_vars.items() if var.get() == 1]

        if not x_axis_var.get():
            messagebox.showerror("错误", "请选择一个横坐标")
            return

        if not selected_columns:
            messagebox.showerror("错误", "请选择至少一个纵坐标")
            return

        draw_plot(x_axis_var.get(), selected_columns)


    # def show_histogram_window():
    #     """绘制柱状图，用于查看某列的数据分布"""
    #     nonlocal current_df
    #
    #     if current_df is None:
    #         messagebox.showerror("错误", "请先选择并加载数据文件")
    #         return
    #
    #     selected_columns = [col for col, var in y_axis_vars.items() if var.get() == 1]
    #     if not selected_columns:
    #         messagebox.showerror("错误", "请选择至少一个纵坐标参数用于绘制柱状图")
    #         return
    #
    #     # 仅使用第一个被选中的列绘图（你也可以改成多个）
    #     target_col = selected_columns[0]
    #     cleaned_data = clean_data(current_df[target_col])
    #     cleaned_data = cleaned_data.dropna()
    #
    #     if cleaned_data.empty:
    #         messagebox.showerror("错误", f"列 {target_col} 无有效数值数据")
    #         return
    #
    #     # 创建新窗口
    #     hist_win = tk.Toplevel()
    #     hist_win.title(f"{target_col} - 柱状图")
    #     hist_win.geometry("600x500")
    #
    #     # 创建图表
    #     fig, ax = plt.subplots(figsize=(6, 4), dpi=100)
    #     ax.hist(cleaned_data, bins=20, edgecolor='black')
    #     ax.set_title(f"{target_col} 的分布")
    #     ax.set_xlabel(target_col)
    #     ax.set_ylabel("频数")
    #     ax.grid(True)
    #
    #     canvas = FigureCanvasTkAgg(fig, master=hist_win)
    #     canvas.draw()
    #     canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    def show_custom_chart_window():
        """根据选择的图类型显示选中所有列的图表"""
        nonlocal current_df

        if current_df is None:
            messagebox.showerror("错误", "请先选择并加载数据文件")
            return

        selected_columns = [col for col, var in y_axis_vars.items() if var.get() == 1]
        if not selected_columns:
            messagebox.showerror("错误", "请选择至少一个纵坐标参数")
            return

        chart_type = chart_type_var.get()

        # 检查数据有效性
        valid_columns = []
        for col in selected_columns:
            cleaned = clean_data(current_df[col]).dropna()
            if not cleaned.empty:
                valid_columns.append((col, cleaned))

        if not valid_columns:
            messagebox.showerror("错误", "选中的列均无有效数据")
            return

        # 创建新窗口
        win = tk.Toplevel()
        win.title(f"图表类型：{chart_type}")
        win.geometry("750x550")

        fig, ax = plt.subplots(figsize=(7, 4.5), dpi=100)

        if chart_type == "柱状图":
            for col, data in valid_columns:
                ax.hist(data, bins=20, alpha=0.5, label=col)
            ax.set_ylabel("频数")

        elif chart_type == "箱线图":
            data_list = [data for _, data in valid_columns]
            labels = [col for col, _ in valid_columns]
            ax.boxplot(data_list, labels=labels, patch_artist=True,
                       boxprops=dict(facecolor='lightblue'))
            ax.set_ylabel("值")

        elif chart_type == "折线图":
            for col, data in valid_columns:
                ax.plot(data.values, marker='o', label=col)
            ax.set_ylabel("值")

        elif chart_type == "散点图":
            for col, data in valid_columns:
                ax.scatter(range(len(data)), data, alpha=0.7, label=col)
            ax.set_ylabel("值")

        ax.set_title(f"图类型：{chart_type}")
        ax.set_xlabel("索引")
        ax.grid(True)
        ax.legend()

        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    confirm_button = tk.Button(main_window, text="绘制图像", command=collect_selected_columns, font=("Arial", 12))
    confirm_button.place(x=800, y=470, width=120, height=30)
    # 添加按钮：显示柱状图


    # 图类型选择下拉框
    tk.Label(main_window, text="图类型:", font=("Arial", 12)).place(x=800, y=510)
    chart_type_combobox = ttk.Combobox(main_window, textvariable=chart_type_var,
                                       values=["柱状图", "箱线图", "折线图", "散点图"],
                                       state='readonly', font=("Arial", 11))
    chart_type_combobox.place(x=950, y=470, width=120, height=30)

    draw_chart_button = tk.Button(main_window, text="显示图表", command=show_custom_chart_window, font=("Arial", 12))
    draw_chart_button.place(x=1070, y=470, width=100, height=30)

    fig = plt.Figure(figsize=(6, 6), dpi=100)
    ax = fig.add_subplot(111)
    canvas = FigureCanvasTkAgg(fig, master=main_window)
    canvas.get_tk_widget().place(x=10, y=50, width=760, height=740)

    def draw_plot(x_col, y_columns):
        """绘制选定列的图表，并启用鼠标缩放和拖动"""
        nonlocal current_df, ax, canvas

        if current_df is None:
            messagebox.showerror("错误", "没有数据可以绘制图表")
            return

        if x_col not in current_df.columns:
            messagebox.showerror("错误", f"横坐标列 '{x_col}' 不存在，请重新选择")
            return

        ax.clear()

        # 读取 X 轴数据
        x_data = current_df[x_col].astype(str)

        try:
            if all(x.count(":") >= 2 for x in x_data):
                today = datetime.datetime.today().strftime('%Y-%m-%d')
                x_data = pd.to_datetime(today + " " + x_data)
            else:
                x_data = pd.to_datetime(x_data)
        except Exception:
            x_data = clean_data(x_data)

        if x_data.isnull().all():
            messagebox.showerror("错误", f"横坐标 '{x_col}' 数据无效，请选择正确列")
            return

        # 绘制 Y 轴数据
        y_data_valid = False
        for column in y_columns:
            if column in current_df.columns:
                y_data = clean_data(current_df[column])
                if y_data.isnull().all():
                    continue
                ax.plot(x_data, y_data, label=column)
                y_data_valid = True

        if not y_data_valid:
            messagebox.showerror("错误", "所有选定的纵坐标列都无效，请选择正确的列")
            return

        # 设置 X 轴格式
        ax.xaxis.set_major_formatter(mdates.DateFormatter("%H:%M:%S"))
        ax.xaxis.set_major_locator(mdates.MinuteLocator(interval=10))  # 每10分钟显示一个刻度
        ax.xaxis.set_minor_locator(mdates.MinuteLocator(interval=2))  # 每2分钟一个小刻度

        # 使用 MaxNLocator 控制最大刻度数，避免横坐标太密集
        ax.xaxis.set_major_locator(MaxNLocator(nbins=6))  # 限制最多显示6个主刻度

        ax.tick_params(axis='x', rotation=45)

        ax.set_xlabel(x_col)
        ax.set_ylabel("遥测值")
        ax.legend()
        ax.grid(True)
        #
        # # 保存初始 X/Y 轴范围
        # ax._original_xlim = ax.get_xlim()
        # ax._original_ylim = ax.get_ylim()
        #
        # # 绑定鼠标事件
        # canvas.get_tk_widget().bind("<MouseWheel>", on_scroll)  # 滚轮缩放
        # canvas.get_tk_widget().bind("<ButtonPress-1>", on_press)  # 按下鼠标左键
        # canvas.get_tk_widget().bind("<B1-Motion>", on_drag)  # 拖动鼠标
        # canvas.get_tk_widget().bind("<ButtonRelease-1>", on_release)  # 释放鼠标

        # 重新绘制
        canvas.draw()

        # 数据点悬停显示
        cursor = mplcursors.cursor(ax, hover=True)
        cursor.connect("add", lambda sel: sel.annotation.set_text(
            f"{x_col}={sel.target[0]}\n{sel.artist.get_label()}={sel.target[1]:.2f}"))



    # def on_scroll(event):
    #     """鼠标滚轮缩放功能（以鼠标指针为中心）"""
    #     nonlocal ax, canvas
    #
    #     # 获取当前 X/Y 轴范围
    #     xlim = ax.get_xlim()
    #     ylim = ax.get_ylim()
    #
    #     # **获取鼠标在画布上的位置**
    #     mouse_x, mouse_y = event.x, event.y
    #     mouse_x_data, mouse_y_data = ax.transData.inverted().transform((mouse_x, mouse_y))  # 转换到数据坐标
    #
    #     # **计算缩放因子**
    #     scale_factor = 0.9 if event.delta > 0 else 1.1  # 向上滚动缩小（放大数据），向下滚动放大（缩小数据）
    #
    #     # **计算新的 X/Y 轴范围（基于鼠标位置缩放）**
    #     new_xlim = [mouse_x_data + (x - mouse_x_data) * scale_factor for x in xlim]
    #     new_ylim = [mouse_y_data + (y - mouse_y_data) * scale_factor for y in ylim]
    #
    #     # **限制缩放范围，防止数据消失**
    #     min_xlim, max_xlim = ax._original_xlim
    #     min_ylim, max_ylim = ax._original_ylim
    #
    #     if (new_xlim[1] - new_xlim[0]) < (max_xlim - min_xlim) * 0.01:
    #         return  # 防止缩放到极限
    #
    #     ax.set_xlim(new_xlim)
    #     ax.set_ylim(new_ylim)
    #
    #     # **重新绘制**
    #     canvas.draw()
    #
    # def on_press(event):
    #     """鼠标按下时，记录起始坐标"""
    #     nonlocal ax
    #     ax._drag_start_x = event.x
    #     ax._drag_start_y = event.y
    #
    # def on_drag(event):
    #     """鼠标拖动时移动视图"""
    #     nonlocal ax, canvas
    #
    #     # **计算鼠标位移**
    #     dx = event.x - ax._drag_start_x
    #     dy = event.y - ax._drag_start_y
    #
    #     # **获取当前 X/Y 轴范围**
    #     xlim = ax.get_xlim()
    #     ylim = ax.get_ylim()
    #
    #     # **计算移动量（转换到数据坐标）**
    #     x_shift = -dx * (xlim[1] - xlim[0]) / canvas.get_tk_widget().winfo_width()
    #     y_shift = dy * (ylim[1] - ylim[0]) / canvas.get_tk_widget().winfo_height()
    #
    #     # **计算新的 X/Y 轴范围**
    #     new_xlim = [xlim[0] + x_shift, xlim[1] + x_shift]
    #     new_ylim = [ylim[0] + y_shift, ylim[1] + y_shift]
    #
    #     # **应用新的坐标范围**
    #     ax.set_xlim(new_xlim)
    #     ax.set_ylim(new_ylim)
    #
    #     # **更新起始点**
    #     ax._drag_start_x = event.x
    #     ax._drag_start_y = event.y
    #
    #     # **重新绘制**
    #     canvas.draw()
    #
    # def on_release(event):
    #     """鼠标释放时，停止拖动"""
    #     nonlocal ax
    #     ax._drag_start_x = None
    #     ax._drag_start_y = None


    def on_file_change(*args):
        """当用户从下拉框选择文件时，重新加载数据"""
        nonlocal current_df
        selected_file_name = file_var.get()
        if selected_file_name:
            selected_file = files[file_names.index(selected_file_name)]
            current_df = read_data(selected_file)
            update_column_selection(current_df)

    def get_start_end_time(df, time_column):
        """获取文件中的起始和结束地面时间码"""
        try:
            # 确保时间列是datetime格式
            time_data = pd.to_datetime(df[time_column], errors='coerce')
            # 获取起始时间（最小时间）和结束时间（最大时间）
            start_time = time_data.min()
            end_time = time_data.max()
            return start_time, end_time
        except Exception as e:
            messagebox.showerror("错误", f"获取起止时间失败: {str(e)}")
            return None, None

    def save_images(event=None):
        """保存当前绘图到图片，并生成Word文档"""
        nonlocal current_df, x_axis_var, y_axis_vars

        if current_df is None:
            messagebox.showerror("错误", "没有数据可保存")
            return

        selected_x = x_axis_var.get()
        selected_y_columns = [col for col, var in y_axis_vars.items() if var.get() == 1]

        if not selected_x or not selected_y_columns:
            messagebox.showerror("错误", "请选择横坐标和至少一个纵坐标")
            return

        now = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
        doc_filename = os.path.join(os.getcwd(), f"绘图报告_{now}.docx")

        # 打开现有的文档或创建新的文档
        document = Document()

        x_data = current_df[selected_x]

        try:
            x_data = pd.to_datetime(x_data)
        except:
            x_data = clean_data(x_data)

        if x_data.isnull().all():
            messagebox.showerror("错误", "横坐标数据无效")
            return

        # 清空图像
        ax.clear()

        # 绘制所有选中的 Y 轴曲线
        y_data_valid = False
        for y_col in selected_y_columns:
            y_data = clean_data(current_df[y_col])
            if y_data.isnull().all():
                continue
            ax.plot(x_data, y_data, label=y_col)
            y_data_valid = True

        if not y_data_valid:
            messagebox.showerror("错误", "所有选定的纵坐标列都无效，请选择正确的列")
            return

        ax.set_xlabel(selected_x)
        ax.set_ylabel("遥测值")
        ax.set_title(f"{selected_x} vs 选定参数")
        ax.legend()
        ax.grid(True)

        # 重新绘制
        canvas.draw()

        # 修正：文件名不能包含 `/` 和特殊字符
        valid_filename = f"{now}_{selected_x}.png"
        valid_filename = valid_filename.replace("/", "_").replace("\\", "_").replace(":", "_")

        # 只保存一张图片
        image_filename = os.path.join(os.getcwd(), valid_filename)
        fig.savefig(image_filename, bbox_inches='tight')

        # 添加图片到 Word
        document.add_picture(image_filename, width=Inches(6))
        document.add_paragraph("\n")  # 换行，避免内容混乱

        # 格式化横坐标
        p = document.add_paragraph()
        p.add_run("横坐标: ").bold = True
        p.add_run(selected_x)
        document.add_paragraph("\n")  # 换行

        # 格式化纵坐标
        p = document.add_paragraph()
        p.add_run("纵坐标: ").bold = True
        for y_col in selected_y_columns:
            p.add_run(y_col + "\n")  # 每个参数换行，保证可读性

        # 获取起止时间
        time_column = "地面时间码"  # 假设该列是地面时间码列，根据需要修改
        start_time, end_time = get_start_end_time(current_df, time_column)

        if start_time and end_time:
            # 将起止时间添加到文档中
            document.add_paragraph(f"起始时间: {start_time.strftime('%Y-%m-%d %H:%M:%S')}")
            document.add_paragraph(f"结束时间: {end_time.strftime('%Y-%m-%d %H:%M:%S')}")

        # 保存并返回文档路径
        document.save(doc_filename)
        return doc_filename

    def save_value(doc_filename):
        """保存选定的参数的最大值和最小值"""
        nonlocal current_df, x_axis_var, y_axis_vars

        # 打开现有的文档
        document = Document(doc_filename)

        # 创建表格
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '序号'
        hdr_cells[1].text = '参数标识'
        hdr_cells[2].text = '参数名称'
        hdr_cells[3].text = '设计值'
        hdr_cells[4].text = '测试值'

        # 获取用户选择的纵坐标参数
        selected_parameters = [col for col, var in y_axis_vars.items() if var.get() == 1]

        if not selected_parameters:
            messagebox.showerror("错误", "没有选择任何参数，请选择至少一个纵坐标参数")
            return

        # 遍历已选参数并统计最大最小值
        for idx, param in enumerate(selected_parameters, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)  # 序号
            row_cells[1].text = f"参数{idx}"  # 参数标识（可以自定义逻辑）
            row_cells[2].text = param  # 参数名称
            row_cells[3].text = "N/A"  # 设计值（这里默认N/A，如有设计值可填入）

            # 确保列存在于 DataFrame
            if param in current_df.columns:
                cleaned_data = clean_data(current_df[param])
                max_value = cleaned_data.max()
                min_value = cleaned_data.min()
                row_cells[4].text = f"最大值: {max_value}, 最小值: {min_value}"
            else:
                row_cells[4].text = "N/A"

            # 设置字体格式
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10.5)  # 六号字体
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 保存文档
        document.save(doc_filename)
        messagebox.showinfo("保存Word", f"Word文档已保存为 {doc_filename}")

    def save_images_and_word(event=None):
        doc_filename = save_images()  # 保存图片并返回文档路径
        save_value(doc_filename)  # 在同一文档中追加表格信息

    def add_to_config():
        """将当前选择的 X 轴和 Y 轴存入 peizhi.json，并且为配置命名"""
        # 获取用户输入的配置名称
        config_name = config_name_entry.get()
        if not config_name:
            messagebox.showerror("错误", "请输入配置名称")
            return

        selected_x = x_axis_var.get()
        selected_y_columns = [col for col, var in y_axis_vars.items() if var.get() == 1]

        if not selected_x or not selected_y_columns:
            messagebox.showerror("错误", "请选择横坐标和至少一个纵坐标")
            return

        config_filename = "peizhi.json"
        config_data = {}

        # 读取现有的 JSON 配置文件（如果存在）
        if os.path.exists(config_filename):
            with open(config_filename, "r", encoding="utf-8") as f:
                try:
                    config_data = json.load(f)
                except json.JSONDecodeError:
                    config_data = {}

        selected_file_name = file_var.get()
        if not selected_file_name:
            messagebox.showerror("错误", "未选择文件，无法添加配置")
            return

        # 检查当前配置是否已经存在（相同的 x 轴和 y 轴）
        if selected_file_name in config_data:
            for existing_config in config_data[selected_file_name]:
                if existing_config["x_axis"] == selected_x and set(existing_config["y_axes"]) == set(
                        selected_y_columns):
                    # 如果相同，则提示并返回已保存的配置名称
                    messagebox.showerror("错误", f"配置已存在：{existing_config['name']}")
                    return

        # 为当前文件添加配置
        new_config = {
            "name": config_name,
            "x_axis": selected_x,
            "y_axes": selected_y_columns
        }

        # 如果该文件已有配置，则将新的配置追加进去
        if selected_file_name in config_data:
            config_data[selected_file_name].append(new_config)
        else:
            # 如果文件是第一次配置，创建新的键值对
            config_data[selected_file_name] = [new_config]

        # 保存更新后的配置到 peizhi.json
        with open(config_filename, "w", encoding="utf-8") as f:
            json.dump(config_data, f, ensure_ascii=False, indent=4)

        messagebox.showinfo("成功", f"配置 '{config_name}' 已添加到 {selected_file_name} 的配置中！")
        load_config_for_file()

    def load_config_for_file():
        """每次选择文件后，检查 peizhi.json 中是否已有配置并加载配置按钮"""
        selected_file_name = file_var.get()
        if not selected_file_name:
            return

        config_filename = "peizhi.json"
        if not os.path.exists(config_filename):
            return

        # 读取 JSON 配置文件
        with open(config_filename, "r", encoding="utf-8") as f:
            try:
                config_data = json.load(f)
            except json.JSONDecodeError:
                config_data = {}

        # 如果文件在配置中存在
        if selected_file_name in config_data:
            # 获取文件的配置列表
            file_configs = config_data[selected_file_name]

            # 清空配置按钮区域
            for widget in config_button_frame.winfo_children():
                widget.destroy()

            # **创建一个带有黑色边框的Frame来容纳Canvas和Scrollbar**
            frame_with_border = tk.Frame(config_button_frame, relief="solid", bd=0.5, padx=5, pady=5)  # 添加黑色边框
            frame_with_border.pack(fill="both", expand=True)

            # 创建Canvas
            config_button_canvas = tk.Canvas(frame_with_border)
            config_scrollbar = tk.Scrollbar(frame_with_border, orient="vertical", command=config_button_canvas.yview)

            # 绑定Scrollbar与Canvas的同步滚动
            config_button_canvas.configure(yscrollcommand=config_scrollbar.set)

            # 创建内嵌Frame用于存放Checkbuttons
            frame_inside_canvas = tk.Frame(config_button_canvas)

            # 将Frame放入Canvas内部
            window_inside_canvas = config_button_canvas.create_window((0, 0), window=frame_inside_canvas, anchor="nw")

            def on_mouse_wheel(event):
                """鼠标滚动时滚动 Canvas"""
                if frame_with_border.winfo_containing(event.x_root, event.y_root):
                    config_button_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

            # **绑定鼠标滚轮事件**
            config_button_canvas.bind("<Enter>",
                                      lambda e: config_button_canvas.bind_all("<MouseWheel>", on_mouse_wheel))
            config_button_canvas.bind("<Leave>", lambda e: config_button_canvas.unbind_all("<MouseWheel>"))

            frame_inside_canvas.bind(
                "<Configure>",
                lambda e: config_button_canvas.configure(
                    scrollregion=config_button_canvas.bbox("all")
                )
            )

            # 设置按钮的统一宽度
            button_width = 20  # 设置按钮的宽度，这里您可以手动设置需要的宽度

            # **在内部 Frame 中添加配置按钮**
            for idx, config in enumerate(file_configs):
                config_name = config["name"]
                config_button = tk.Button(frame_inside_canvas, text=config_name,
                                          width=button_width,  # 设置按钮宽度
                                          command=lambda config=config: on_config_button_click(config))
                config_button.grid(row=idx // 2, column=idx % 2, padx=5, pady=5, sticky="w")  # 两列显示

            # **将Canvas和Scrollbar放入容器中**
            config_button_canvas.pack(side="left", fill="both", expand=True)
            config_scrollbar.pack(side="right", fill="y")  # 滚动条可见且可拖动

    def on_config_button_click(config):
        """点击配置按钮时，更新配置名称文本框和纵坐标选择框"""
        # 更新配置名称文本框
        config_name_entry.delete(0, tk.END)
        config_name_entry.insert(0, config["name"])

        # 获取配置的 y_axes 列表
        selected_y_columns = config["y_axes"]

        # 更新纵坐标多选框的选中状态
        for col, var in y_axis_vars.items():
            # 如果该列在配置中，则选中复选框
            if col in selected_y_columns:
                var.set(1)  # 选中该复选框
            else:
                var.set(0)  # 取消选中

    def on_file_change(*args):
        """当用户从下拉框选择文件时，加载该文件的配置"""
        load_config_for_file()

    # **修改：文件选择框的 trace，添加回调事件**
    file_var.trace('w', on_file_change)

    # **创建可滚动的 Frame 以存放配置按钮**
    config_button_frame = tk.Frame(main_window)
    config_button_frame.place(x=800, y=550, width=350, height=230)  # 配置按钮的展示区域

    # **在 UI 上添加一个文本框供用户输入配置名称**
    tk.Label(main_window, text="配置名称:", font=("Arial", 12)).place(x=800, y=510)  # 标签
    config_name_entry = tk.Entry(main_window, font=("Arial", 12))
    config_name_entry.place(x=880, y=510, width=150, height=25)  # 文本框

    # **添加 “添加配置” 按钮**
    add_config_button = tk.Button(main_window, text="添加配置", command=add_to_config, font=("Arial", 12))
    add_config_button.place(x=1040, y=510, width=100, height=30)  # 放置按钮

    menu = tk.Menu(main_window, tearoff=0)
    menu.add_command(label="保存图片和生成Word文档", command=save_images_and_word)
    canvas.get_tk_widget().bind("<Button-3>", lambda event: menu.tk_popup(event.x_root, event.y_root))

    main_window.mainloop()


def start_program():
    root = tk.Tk()
    root.title("文件读取程序")
    root.geometry("400x200")

    def open_folder_selection():
        folder_path = select_folder()
        if folder_path:
            root.destroy()
            create_main_interface(folder_path)

    button = tk.Button(root, text="选择文件夹", command=open_folder_selection, font=("Arial", 12))
    button.pack(expand=True)

    root.mainloop()


if __name__ == '__main__':
    start_program()